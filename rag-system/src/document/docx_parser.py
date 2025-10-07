from docx import Document
from typing import List, Optional
from datetime import datetime
from pathlib import Path
import logging

from .parser import DocumentParser, DocumentMetadata, ParsedDocument

logger = logging.getLogger(__name__)

class DocxParser(DocumentParser):
    """Word文档解析器"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def __init__(self):
        super().__init__()
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否可以解析Word文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            bool: 是否可以解析
        """
        try:
            path = Path(file_path)
            return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
        except Exception as e:
            self.logger.error(f"检查文件类型失败: {e}")
            return False
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文档
        
        Args:
            file_path: Word文件路径
            
        Returns:
            ParsedDocument: 解析后的文档数据
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持或损坏
            Exception: 解析过程中的其他错误
        """
        self.validate_file(file_path)
        
        if not self.can_parse(file_path):
            raise ValueError(f"不支持的文件格式: {file_path}")
        
        try:
            # 打开Word文档
            doc = Document(file_path)
            
            # 提取文本内容
            content = self._extract_text(doc)
            
            # 提取元数据
            metadata = self._extract_metadata_from_doc(doc, file_path)
            
            # 清理文本
            content = self.clean_text(content)
            
            # 检测语言
            if not metadata.language:
                metadata.language = self.detect_language(content)
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                file_path=file_path,
                file_type='docx'
            )
            
        except Exception as e:
            self.logger.error(f"解析Word文件失败 {file_path}: {e}")
            raise Exception(f"Word解析失败: {e}")
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """提取Word文档元数据
        
        Args:
            file_path: Word文件路径
            
        Returns:
            DocumentMetadata: 文档元数据
        """
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            return self._extract_metadata_from_doc(doc, file_path)
        except Exception as e:
            self.logger.error(f"提取Word元数据失败 {file_path}: {e}")
            # 返回基本元数据
            file_info = self.get_file_info(file_path)
            return DocumentMetadata(
                title=file_info['file_name'],
                file_size=file_info['file_size'],
                creation_date=file_info['creation_time'],
                modification_date=file_info['modification_time']
            )
    
    def _extract_text(self, doc: Document) -> str:
        """从Word文档中提取文本
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            str: 提取的文本内容
        """
        text_content = []
        
        try:
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_content.append(table_text)
            
        except Exception as e:
            self.logger.warning(f"提取Word文本失败: {e}")
        
        return '\n\n'.join(text_content)
    
    def _extract_table_text(self, table) -> str:
        """从表格中提取文本
        
        Args:
            table: python-docx表格对象
            
        Returns:
            str: 表格文本内容
        """
        table_text = []
        
        try:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                
                if row_text:
                    table_text.append(' | '.join(row_text))
        
        except Exception as e:
            self.logger.warning(f"提取表格文本失败: {e}")
        
        return '\n'.join(table_text)
    
    def _extract_metadata_from_doc(self, doc: Document, file_path: str) -> DocumentMetadata:
        """从Word文档对象中提取元数据
        
        Args:
            doc: python-docx文档对象
            file_path: 文件路径
            
        Returns:
            DocumentMetadata: 文档元数据
        """
        try:
            # 获取文档属性
            core_props = doc.core_properties
            
            # 获取文件信息
            file_info = self.get_file_info(file_path)
            
            # 提取关键词
            keywords = None
            if core_props.keywords:
                keywords = [kw.strip() for kw in core_props.keywords.split(',') if kw.strip()]
            
            # 计算页面数（估算）
            page_count = self._estimate_page_count(doc)
            
            return DocumentMetadata(
                title=core_props.title or file_info['file_name'],
                author=core_props.author,
                creation_date=core_props.created or file_info['creation_time'],
                modification_date=core_props.modified or file_info['modification_time'],
                page_count=page_count,
                file_size=file_info['file_size'],
                keywords=keywords,
                subject=core_props.subject
            )
            
        except Exception as e:
            self.logger.warning(f"提取Word元数据失败: {e}")
            # 返回基本元数据
            file_info = self.get_file_info(file_path)
            return DocumentMetadata(
                title=file_info['file_name'],
                file_size=file_info['file_size'],
                creation_date=file_info['creation_time'],
                modification_date=file_info['modification_time']
            )
    
    def _estimate_page_count(self, doc: Document) -> Optional[int]:
        """估算Word文档页面数
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            Optional[int]: 估算的页面数
        """
        try:
            # 简单估算：每页约500个字符
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            if total_chars > 0:
                return max(1, total_chars // 500)
        except Exception as e:
            self.logger.warning(f"估算页面数失败: {e}")
        
        return None
    
    def extract_paragraphs(self, file_path: str) -> List[str]:
        """提取文档段落
        
        Args:
            file_path: Word文件路径
            
        Returns:
            List[str]: 段落文本列表
        """
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = self.clean_text(paragraph.text)
                if text:
                    paragraphs.append(text)
            
            return paragraphs
            
        except Exception as e:
            self.logger.error(f"提取Word段落失败 {file_path}: {e}")
            raise Exception(f"提取Word段落失败: {e}")
    
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """提取文档表格
        
        Args:
            file_path: Word文件路径
            
        Returns:
            List[List[List[str]]]: 表格数据，格式为[表格][行][列]
        """
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = self.clean_text(cell.text)
                        row_data.append(cell_text)
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return tables_data
            
        except Exception as e:
            self.logger.error(f"提取Word表格失败 {file_path}: {e}")
            raise Exception(f"提取Word表格失败: {e}")
    
    def get_paragraph_count(self, file_path: str) -> int:
        """获取段落数
        
        Args:
            file_path: Word文件路径
            
        Returns:
            int: 段落数
        """
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            return len([p for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            self.logger.error(f"获取Word段落数失败 {file_path}: {e}")
            return 0