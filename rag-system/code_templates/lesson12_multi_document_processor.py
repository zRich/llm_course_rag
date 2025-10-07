#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Lesson12 多文档处理器实现模板
解决多格式文档支持、文档源管理和并发处理缺失问题

功能特性：
1. 多格式文档支持（PDF、Word、Excel、TXT、Markdown等）
2. 文档源管理和版本控制
3. 并发处理和批量操作
4. 文档元数据提取和管理
5. 错误处理和恢复机制
"""

import logging
import asyncio
import aiofiles
from typing import List, Dict, Optional, Union, AsyncGenerator, Tuple
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime
import hashlib
import json
import mimetypes
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from queue import Queue

# 文档处理库
import fitz  # PyMuPDF for PDF
import docx  # python-docx for Word
import pandas as pd  # for Excel
import markdown  # for Markdown
from bs4 import BeautifulSoup  # for HTML
import chardet  # for encoding detection

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """文档元数据"""
    file_path: str
    file_name: str
    file_size: int
    file_type: str
    mime_type: str
    encoding: str
    created_time: datetime
    modified_time: datetime
    processed_time: datetime
    content_hash: str
    page_count: int = 0
    word_count: int = 0
    char_count: int = 0
    language: str = "unknown"
    author: str = ""
    title: str = ""
    subject: str = ""
    keywords: List[str] = field(default_factory=list)
    custom_fields: Dict = field(default_factory=dict)

@dataclass
class ProcessedDocument:
    """处理后的文档"""
    doc_id: str
    content: str
    metadata: DocumentMetadata
    sections: List[Dict] = field(default_factory=list)
    tables: List[Dict] = field(default_factory=list)
    images: List[Dict] = field(default_factory=list)
    links: List[str] = field(default_factory=list)
    processing_errors: List[str] = field(default_factory=list)
    processing_time: float = 0.0

class DocumentProcessor:
    """文档处理器基类"""
    
    def __init__(self):
        self.supported_extensions = set()
        self.supported_mime_types = set()
    
    def can_process(self, file_path: str, mime_type: str = None) -> bool:
        """检查是否支持处理该文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions or mime_type in self.supported_mime_types
    
    def process(self, file_path: str) -> ProcessedDocument:
        """处理文档（需要子类实现）"""
        raise NotImplementedError
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """提取基础元数据"""
        path = Path(file_path)
        stat = path.stat()
        
        # 检测文件编码
        encoding = self._detect_encoding(file_path)
        
        # 生成内容哈希
        content_hash = self._calculate_file_hash(file_path)
        
        return DocumentMetadata(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_size=stat.st_size,
            file_type=path.suffix.lower(),
            mime_type=mimetypes.guess_type(file_path)[0] or "unknown",
            encoding=encoding,
            created_time=datetime.fromtimestamp(stat.st_ctime),
            modified_time=datetime.fromtimestamp(stat.st_mtime),
            processed_time=datetime.now(),
            content_hash=content_hash
        )
    
    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # 读取前10KB
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except Exception:
            return 'utf-8'
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception:
            return ""

class PDFProcessor(DocumentProcessor):
    """PDF文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.pdf'}
        self.supported_mime_types = {'application/pdf'}
    
    def process(self, file_path: str) -> ProcessedDocument:
        """处理PDF文档"""
        start_time = datetime.now()
        doc_id = self._generate_doc_id(file_path)
        
        try:
            # 提取元数据
            metadata = self.extract_metadata(file_path)
            
            # 打开PDF文档
            pdf_doc = fitz.open(file_path)
            
            # 提取PDF特定元数据
            pdf_metadata = pdf_doc.metadata
            metadata.title = pdf_metadata.get('title', '')
            metadata.author = pdf_metadata.get('author', '')
            metadata.subject = pdf_metadata.get('subject', '')
            metadata.page_count = len(pdf_doc)
            
            # 提取内容
            content_parts = []
            sections = []
            tables = []
            images = []
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                
                # 提取文本
                page_text = page.get_text()
                if page_text.strip():
                    content_parts.append(page_text)
                    sections.append({
                        'type': 'text',
                        'page': page_num + 1,
                        'content': page_text,
                        'bbox': None
                    })
                
                # 提取表格
                try:
                    page_tables = page.find_tables()
                    for i, table in enumerate(page_tables):
                        table_data = table.extract()
                        if table_data:
                            tables.append({
                                'page': page_num + 1,
                                'table_id': f"table_{page_num}_{i}",
                                'data': table_data,
                                'bbox': table.bbox
                            })
                except Exception as e:
                    logger.warning(f"表格提取失败 (页面 {page_num + 1}): {e}")
                
                # 提取图像信息
                try:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        images.append({
                            'page': page_num + 1,
                            'image_id': f"img_{page_num}_{img_index}",
                            'xref': img[0],
                            'bbox': None
                        })
                except Exception as e:
                    logger.warning(f"图像提取失败 (页面 {page_num + 1}): {e}")
            
            pdf_doc.close()
            
            # 合并内容
            full_content = '\n\n'.join(content_parts)
            metadata.word_count = len(full_content.split())
            metadata.char_count = len(full_content)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                doc_id=doc_id,
                content=full_content,
                metadata=metadata,
                sections=sections,
                tables=tables,
                images=images,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"PDF处理失败 {file_path}: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return ProcessedDocument(
                doc_id=doc_id,
                content="",
                metadata=self.extract_metadata(file_path),
                processing_errors=[str(e)],
                processing_time=processing_time
            )
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]

class WordProcessor(DocumentProcessor):
    """Word文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.docx', '.doc'}
        self.supported_mime_types = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        }
    
    def process(self, file_path: str) -> ProcessedDocument:
        """处理Word文档"""
        start_time = datetime.now()
        doc_id = self._generate_doc_id(file_path)
        
        try:
            # 提取元数据
            metadata = self.extract_metadata(file_path)
            
            # 打开Word文档
            doc = docx.Document(file_path)
            
            # 提取Word特定元数据
            core_props = doc.core_properties
            metadata.title = core_props.title or ''
            metadata.author = core_props.author or ''
            metadata.subject = core_props.subject or ''
            
            # 提取内容
            content_parts = []
            sections = []
            tables = []
            
            # 提取段落
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
                    sections.append({
                        'type': 'paragraph',
                        'index': i,
                        'content': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None
                    })
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    'table_id': f"table_{i}",
                    'data': table_data,
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0
                })
            
            # 合并内容
            full_content = '\n\n'.join(content_parts)
            metadata.word_count = len(full_content.split())
            metadata.char_count = len(full_content)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                doc_id=doc_id,
                content=full_content,
                metadata=metadata,
                sections=sections,
                tables=tables,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Word处理失败 {file_path}: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return ProcessedDocument(
                doc_id=doc_id,
                content="",
                metadata=self.extract_metadata(file_path),
                processing_errors=[str(e)],
                processing_time=processing_time
            )
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]

class ExcelProcessor(DocumentProcessor):
    """Excel文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.xlsx', '.xls', '.csv'}
        self.supported_mime_types = {
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'text/csv'
        }
    
    def process(self, file_path: str) -> ProcessedDocument:
        """处理Excel文档"""
        start_time = datetime.now()
        doc_id = self._generate_doc_id(file_path)
        
        try:
            # 提取元数据
            metadata = self.extract_metadata(file_path)
            
            # 读取Excel文件
            if file_path.endswith('.csv'):
                df_dict = {'Sheet1': pd.read_csv(file_path)}
            else:
                df_dict = pd.read_excel(file_path, sheet_name=None)
            
            # 提取内容
            content_parts = []
            tables = []
            
            for sheet_name, df in df_dict.items():
                # 转换为文本
                sheet_text = f"工作表: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                content_parts.append(sheet_text)
                
                # 保存表格数据
                tables.append({
                    'sheet_name': sheet_name,
                    'data': df.values.tolist(),
                    'columns': df.columns.tolist(),
                    'rows': len(df),
                    'cols': len(df.columns)
                })
            
            # 合并内容
            full_content = '\n\n'.join(content_parts)
            metadata.word_count = len(full_content.split())
            metadata.char_count = len(full_content)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                doc_id=doc_id,
                content=full_content,
                metadata=metadata,
                tables=tables,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Excel处理失败 {file_path}: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return ProcessedDocument(
                doc_id=doc_id,
                content="",
                metadata=self.extract_metadata(file_path),
                processing_errors=[str(e)],
                processing_time=processing_time
            )
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]

class TextProcessor(DocumentProcessor):
    """文本文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.txt', '.md', '.markdown', '.html', '.htm'}
        self.supported_mime_types = {
            'text/plain',
            'text/markdown',
            'text/html'
        }
    
    def process(self, file_path: str) -> ProcessedDocument:
        """处理文本文档"""
        start_time = datetime.now()
        doc_id = self._generate_doc_id(file_path)
        
        try:
            # 提取元数据
            metadata = self.extract_metadata(file_path)
            
            # 读取文件内容
            with open(file_path, 'r', encoding=metadata.encoding) as f:
                raw_content = f.read()
            
            # 根据文件类型处理内容
            if file_path.endswith(('.md', '.markdown')):
                # Markdown处理
                html_content = markdown.markdown(raw_content)
                soup = BeautifulSoup(html_content, 'html.parser')
                content = soup.get_text()
            elif file_path.endswith(('.html', '.htm')):
                # HTML处理
                soup = BeautifulSoup(raw_content, 'html.parser')
                content = soup.get_text()
                # 提取链接
                links = [a.get('href') for a in soup.find_all('a', href=True)]
            else:
                # 纯文本
                content = raw_content
                links = []
            
            metadata.word_count = len(content.split())
            metadata.char_count = len(content)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                doc_id=doc_id,
                content=content,
                metadata=metadata,
                links=links if 'links' in locals() else [],
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"文本处理失败 {file_path}: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()
            return ProcessedDocument(
                doc_id=doc_id,
                content="",
                metadata=self.extract_metadata(file_path),
                processing_errors=[str(e)],
                processing_time=processing_time
            )
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]

class MultiDocumentProcessor:
    """多文档处理器管理器"""
    
    def __init__(self, max_workers: int = 4):
        self.processors = {
            'pdf': PDFProcessor(),
            'word': WordProcessor(),
            'excel': ExcelProcessor(),
            'text': TextProcessor()
        }
        self.max_workers = max_workers
        self.processing_queue = Queue()
        self.results_queue = Queue()
        self._lock = threading.Lock()
    
    def get_processor(self, file_path: str) -> Optional[DocumentProcessor]:
        """获取适合的处理器"""
        mime_type = mimetypes.guess_type(file_path)[0]
        
        for processor in self.processors.values():
            if processor.can_process(file_path, mime_type):
                return processor
        
        return None
    
    def process_single(self, file_path: str) -> ProcessedDocument:
        """处理单个文档"""
        processor = self.get_processor(file_path)
        if not processor:
            raise ValueError(f"不支持的文件类型: {file_path}")
        
        return processor.process(file_path)
    
    def process_batch(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """批量处理文档"""
        results = []
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # 提交任务
            future_to_path = {
                executor.submit(self.process_single, path): path 
                for path in file_paths
            }
            
            # 收集结果
            for future in as_completed(future_to_path):
                file_path = future_to_path[future]
                try:
                    result = future.result()
                    results.append(result)
                    logger.info(f"处理完成: {file_path}")
                except Exception as e:
                    logger.error(f"处理失败 {file_path}: {e}")
                    # 创建错误结果
                    error_result = ProcessedDocument(
                        doc_id=hashlib.md5(file_path.encode()).hexdigest()[:16],
                        content="",
                        metadata=DocumentMetadata(
                            file_path=file_path,
                            file_name=Path(file_path).name,
                            file_size=0,
                            file_type="unknown",
                            mime_type="unknown",
                            encoding="unknown",
                            created_time=datetime.now(),
                            modified_time=datetime.now(),
                            processed_time=datetime.now(),
                            content_hash=""
                        ),
                        processing_errors=[str(e)]
                    )
                    results.append(error_result)
        
        return results
    
    async def process_batch_async(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """异步批量处理文档"""
        loop = asyncio.get_event_loop()
        
        # 创建任务
        tasks = []
        for file_path in file_paths:
            task = loop.run_in_executor(
                None, self.process_single, file_path
            )
            tasks.append(task)
        
        # 等待所有任务完成
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 处理异常结果
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"异步处理失败 {file_paths[i]}: {result}")
                error_result = ProcessedDocument(
                    doc_id=hashlib.md5(file_paths[i].encode()).hexdigest()[:16],
                    content="",
                    metadata=DocumentMetadata(
                        file_path=file_paths[i],
                        file_name=Path(file_paths[i]).name,
                        file_size=0,
                        file_type="unknown",
                        mime_type="unknown",
                        encoding="unknown",
                        created_time=datetime.now(),
                        modified_time=datetime.now(),
                        processed_time=datetime.now(),
                        content_hash=""
                    ),
                    processing_errors=[str(result)]
                )
                processed_results.append(error_result)
            else:
                processed_results.append(result)
        
        return processed_results
    
    def scan_directory(self, directory: str, recursive: bool = True) -> List[str]:
        """扫描目录获取支持的文件"""
        supported_files = []
        directory_path = Path(directory)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"目录不存在: {directory}")
        
        # 获取所有支持的扩展名
        supported_extensions = set()
        for processor in self.processors.values():
            supported_extensions.update(processor.supported_extensions)
        
        # 扫描文件
        pattern = "**/*" if recursive else "*"
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                supported_files.append(str(file_path))
        
        logger.info(f"在目录 {directory} 中找到 {len(supported_files)} 个支持的文件")
        return supported_files
    
    def get_processing_stats(self, results: List[ProcessedDocument]) -> Dict:
        """获取处理统计信息"""
        total_files = len(results)
        successful = len([r for r in results if not r.processing_errors])
        failed = total_files - successful
        
        total_processing_time = sum(r.processing_time for r in results)
        total_content_length = sum(len(r.content) for r in results)
        total_word_count = sum(r.metadata.word_count for r in results)
        
        file_types = {}
        for result in results:
            file_type = result.metadata.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_files': total_files,
            'successful': successful,
            'failed': failed,
            'success_rate': successful / total_files if total_files > 0 else 0,
            'total_processing_time': total_processing_time,
            'avg_processing_time': total_processing_time / total_files if total_files > 0 else 0,
            'total_content_length': total_content_length,
            'total_word_count': total_word_count,
            'file_types': file_types
        }

def main():
    """示例用法"""
    # 初始化多文档处理器
    processor = MultiDocumentProcessor(max_workers=4)
    
    # 示例文件路径
    file_paths = [
        "example.pdf",
        "example.docx",
        "example.xlsx",
        "example.txt"
    ]
    
    # 过滤存在的文件
    existing_files = [f for f in file_paths if Path(f).exists()]
    
    if existing_files:
        print(f"处理 {len(existing_files)} 个文件...")
        
        # 批量处理
        results = processor.process_batch(existing_files)
        
        # 显示结果
        for result in results:
            print(f"\n文档ID: {result.doc_id}")
            print(f"文件: {result.metadata.file_name}")
            print(f"类型: {result.metadata.file_type}")
            print(f"大小: {result.metadata.file_size} bytes")
            print(f"字数: {result.metadata.word_count}")
            print(f"处理时间: {result.processing_time:.2f}s")
            if result.processing_errors:
                print(f"错误: {result.processing_errors}")
            else:
                print(f"内容预览: {result.content[:100]}...")
        
        # 显示统计信息
        stats = processor.get_processing_stats(results)
        print(f"\n处理统计:")
        print(f"总文件数: {stats['total_files']}")
        print(f"成功: {stats['successful']}")
        print(f"失败: {stats['failed']}")
        print(f"成功率: {stats['success_rate']:.2%}")
        print(f"总处理时间: {stats['total_processing_time']:.2f}s")
        print(f"平均处理时间: {stats['avg_processing_time']:.2f}s")
        print(f"文件类型分布: {stats['file_types']}")
    else:
        print("没有找到示例文件")

if __name__ == "__main__":
    main()